import sys
import io
import re
import json
import time
import os
import threading
from datetime import datetime, timedelta
from pathlib import Path
from functools import wraps
from collections import defaultdict
from flask import Flask, request, jsonify, send_from_directory, session, redirect, url_for, Response, stream_with_context
from flask_cors import CORS
import pdfplumber
from openai import OpenAI


def load_local_env():
    env_path = Path(__file__).with_name(".env")
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


load_local_env()
import config

app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024
app.secret_key = config.SECRET_KEY
client = None


def print_error(msg):
    print("\n" + "!" * 60)
    print("❌ 配置错误")
    print("!" * 60)
    print(msg)
    print("!" * 60 + "\n")


def check_config():
    if not getattr(config, "API_KEY", "") or config.API_KEY == "PASTE_KEY_HERE":
        print_error("请先设置环境变量 OPENROUTER_API_KEY。不要把 API Key 写死在代码里或粘贴到聊天窗口。")
        sys.exit(1)
    if not config.API_KEY.startswith("sk-or-v1-"):
        print_error("API Key 格式不对。OpenRouter Key 应该以 sk-or-v1- 开头。")
        sys.exit(1)
    if getattr(config, "API_BASE", "") != "https://openrouter.ai/api/v1":
        print_error("API_BASE 应为 https://openrouter.ai/api/v1")
        sys.exit(1)


def read_file(path):
    p = Path(path)
    return p.read_text(encoding="utf-8") if p.exists() else ""


def current_date():
    return datetime.now().strftime("%Y-%m-%d")


def login_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        if not config.APP_PASSWORD or session.get("authenticated"):
            return func(*args, **kwargs)
        if request.path.startswith("/api/"):
            return jsonify({"error": "请先登录后再使用。"}), 401
        return redirect(url_for("login"))
    return wrapper


@app.route("/login", methods=["GET", "POST"])
def login():
    if not config.APP_PASSWORD:
        return redirect(url_for("index"))
    error = ""
    if request.method == "POST":
        password = request.form.get("password", "")
        if password == config.APP_PASSWORD:
            session["authenticated"] = True
            return redirect(url_for("index"))
        error = "密码不正确"
    return f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>登录 · CEO 面试副驾驶</title>
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ min-height: 100vh; display: flex; align-items: center; justify-content: center; font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", sans-serif; background: linear-gradient(135deg, #f5f7fb 0%, #eef4fb 100%); color: #2c3e50; padding: 24px; }}
    .login-card {{ width: 100%; max-width: 420px; background: #fff; border-radius: 18px; padding: 32px; box-shadow: 0 12px 40px rgba(44,62,80,0.12); border: 1px solid #edf1f5; }}
    h1 {{ font-size: 24px; margin-bottom: 8px; }}
    p {{ color: #718096; font-size: 14px; margin-bottom: 22px; line-height: 1.7; }}
    label {{ display: block; font-size: 13px; margin-bottom: 8px; color: #4a5568; font-weight: 600; }}
    input {{ width: 100%; border: 1px solid #d9e2ec; border-radius: 10px; padding: 12px 14px; font-size: 15px; outline: none; }}
    input:focus {{ border-color: #3498db; box-shadow: 0 0 0 3px rgba(52,152,219,0.12); }}
    button {{ width: 100%; margin-top: 16px; border: none; background: #3498db; color: #fff; border-radius: 10px; padding: 12px; font-size: 15px; font-weight: 700; cursor: pointer; }}
    button:hover {{ background: #2980b9; }}
    .error {{ margin-top: 14px; color: #c0392b; font-size: 13px; }}
  </style>
</head>
<body>
  <form class="login-card" method="post">
    <h1>🎯 CEO 面试副驾驶</h1>
    <p>该系统包含候选人隐私和公司面试标准，请输入访问密码。</p>
    <label>访问密码</label>
    <input type="password" name="password" autofocus>
    <button type="submit">进入系统</button>
    <div class="error">{error}</div>
  </form>
</body>
</html>
"""


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


NAME_NOISE_RE = re.compile(r"简历|候选人|应聘|求职|岗位|职位|负责人|主管|经理|总监|专员|招聘|面试")


def extract_person_name(raw):
    s = (raw or "").strip()
    if not s:
        return ""
    s = re.sub(r"\.(pdf|docx?|txt|md|json)$", "", s, flags=re.IGNORECASE)
    prev = None
    while prev != s:
        prev = s
        s = re.sub(r"[【（(\[][^【】（）()\[\]]*[】）)\]]", " ", s)
    parts = re.split(r"[\s_\-—·｜|,，、/\\]+", s)
    cand = [p for p in parts if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", p) and not NAME_NOISE_RE.search(p)]
    if cand:
        return cand[0]
    m = re.findall(r"[\u4e00-\u9fa5]{2,4}", NAME_NOISE_RE.sub(" ", s))
    return m[0] if m else ""


def clean_person_name(name):
    n = (name or "").strip()
    if not n:
        return ""
    if re.search(r"[【（(\[]", n) or re.search(r"\d", n) or NAME_NOISE_RE.search(n):
        p = extract_person_name(n)
        if p:
            return p
    return n


def extract_candidate_name(markdown):
    patterns = [
        r"^#\s*([^\n#：:|]+)",
        r"姓名[：:]\s*([^\n，,；;\s]+)",
        r"候选人姓名[：:]\s*([^\n，,；;\s]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, markdown or "", re.MULTILINE)
        if match:
            name = match.group(1).strip()
            if name and name not in {"候选人", "未提供", "（未提供）"}:
                return clean_person_name(name) or name
    return "候选人"


def jd_prompt(raw_description):
    return f"""
你是一位资深招聘负责人。请把用户输入的岗位描述整理成一份正式、清晰、可用于候选人适配度评估的 JD。

## 当前日期
{current_date()}

## 用户输入的岗位描述
{raw_description}

## 输出要求
1. 不要编造公司未提供的硬性条件；如果信息缺失，写“待补充”。
2. 重点服务于后续人才评估，明确职责、能力要求、加分项、风险点和一票否决。
3. 用 Markdown 输出。

## 输出格式
# 正式岗位描述

## 岗位定位

## 核心职责

## 必备能力

## 加分项

## 关键业务场景

## 本岗位特殊权重与一票否决建议

## 用于面试验证的重点
"""


def generate_formal_jd(raw_description):
    return call_llm(jd_prompt(raw_description))


USAGE_LOG_PATH = Path("outputs") / "usage_log.json"


def log_usage(usage):
    if not usage:
        return
    prompt_tokens = getattr(usage, "prompt_tokens", 0)
    completion_tokens = getattr(usage, "completion_tokens", 0)
    total_tokens = getattr(usage, "total_tokens", 0)
    # Claude Opus 4.8 参考价: Input $5/MTok, Output $25/MTok
    cost = prompt_tokens * 0.000005 + completion_tokens * 0.000025
    print(f"\n[Token Usage] 输入: {prompt_tokens} | 输出: {completion_tokens} | 总计: {total_tokens} | 预估费用: ${cost:.4f}\n")
    # 持久化
    try:
        records = []
        if USAGE_LOG_PATH.exists():
            records = json.loads(USAGE_LOG_PATH.read_text(encoding="utf-8"))
            if not isinstance(records, list):
                records = []
        record = {
            "timestamp": datetime.now().isoformat(),
            "date": datetime.now().strftime("%Y-%m-%d"),
            "source": "openrouter",
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
            "total_tokens": total_tokens,
            "cost": round(cost, 6),
            "model": getattr(config, "MODEL", "unknown"),
        }
        records.append(record)
        USAGE_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        tmp = USAGE_LOG_PATH.with_suffix(".json.tmp")
        tmp.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(USAGE_LOG_PATH)
    except Exception as e:
        print(f"[Usage Log Error] {e}")


def get_client():
    global client
    if client is None:
        if not getattr(config, "API_KEY", ""):
            raise RuntimeError("请先设置环境变量 OPENROUTER_API_KEY。")
        client = OpenAI(
            api_key=config.API_KEY,
            base_url=config.API_BASE,
            timeout=getattr(config, "LLM_TIMEOUT", 240),
            max_retries=0,
        )
    return client


def _chat_once(messages):
    response = get_client().chat.completions.create(
        model=config.MODEL,
        max_tokens=config.MAX_TOKENS,
        messages=messages,
        extra_headers={
            "HTTP-Referer": "http://localhost:5000",
            "X-Title": "CEO Interview Copilot",
        },
    )
    log_usage(response.usage)
    choice = response.choices[0]
    return choice.message.content or "", getattr(choice, "finish_reason", None)


def call_llm(prompt):
    last_error = None
    max_continuations = int(getattr(config, "MAX_CONTINUATIONS", 4))
    for attempt in range(2):
        try:
            print(f"[LLM] attempt={attempt + 1} prompt_chars={len(prompt)} model={config.MODEL}")
            messages = [{"role": "user", "content": prompt}]
            content, finish_reason = _chat_once(messages)
            # 自动续写：被截断时让模型从断点继续，拼接成完整内容
            continuations = 0
            while finish_reason == "length" and continuations < max_continuations:
                continuations += 1
                print(f"[LLM] 输出被截断，自动续写 第 {continuations} 次")
                tail = content[-2000:]
                messages = [
                    {"role": "user", "content": prompt},
                    {"role": "assistant", "content": content},
                    {"role": "user", "content": (
                        "上一条回复因长度限制被截断了。请直接从被截断处继续输出剩余内容，"
                        "不要重复已经输出的部分，不要加任何前言或寒暄，不要重写开头。"
                        f"\n\n（已输出内容的结尾片段，供你定位续写位置）：\n{tail}"
                    )},
                ]
                more, finish_reason = _chat_once(messages)
                content += more
            if finish_reason == "length":
                print(f"[LLM WARNING] 续写 {max_continuations} 次后仍被截断 max_tokens={config.MAX_TOKENS}")
                content += (
                    f"\n\n---\n> ⚠️ **内容较长，自动续写 {max_continuations} 次后仍未完成。**"
                    f"\n> 可调高环境变量 `MAX_TOKENS` 或 `MAX_CONTINUATIONS` 后重启服务再试。"
                )
            elif continuations:
                print(f"[LLM] 自动续写完成，共续写 {continuations} 次，总长度 {len(content)} 字符")
            return content
        except Exception as e:
            last_error = e
            print(f"[LLM ERROR] attempt={attempt + 1} type={type(e).__name__} error={str(e)}")
            error_msg = str(e).lower()
            if isinstance(e, json.JSONDecodeError) or "expecting value" in error_msg:
                if attempt == 0:
                    time.sleep(1)
                    continue
                raise RuntimeError("OpenRouter 返回了异常或不完整响应，请稍后重试。如果反复出现，可能是输入内容过长或上游服务临时异常。")
            if "401" in error_msg or "authentication" in error_msg or "api key" in error_msg:
                raise RuntimeError("OpenRouter API Key 无效或未配置，请检查环境变量 OPENROUTER_API_KEY。")
            if "quota" in error_msg or "credit" in error_msg or "billing" in error_msg or "insufficient" in error_msg:
                raise RuntimeError("OpenRouter 账户额度不足或计费异常，请检查 OpenRouter 余额。")
            if "rate" in error_msg or "429" in error_msg:
                raise RuntimeError("调用频率过高，请稍后重试。")
            if "timeout" in error_msg or "connection" in error_msg or "network" in error_msg:
                raise RuntimeError("连接 OpenRouter 失败，请检查网络。")
            break
    raise RuntimeError(f"LLM 调用失败：{str(last_error)}")


def friendly_llm_error(msg):
    m = (msg or "").lower()
    if "401" in m or "authentication" in m or "api key" in m:
        return "OpenRouter API Key 无效或未配置，请检查环境变量 OPENROUTER_API_KEY。"
    if "quota" in m or "credit" in m or "billing" in m or "insufficient" in m:
        return "OpenRouter 账户额度不足或计费异常，请检查 OpenRouter 余额。"
    if "rate" in m or "429" in m:
        return "调用频率过高，请稍后重试。"
    if "timeout" in m or "connection" in m or "network" in m:
        return "连接 OpenRouter 失败，请检查网络后重试。"
    return f"生成失败：{msg}"


def _chat_stream(messages):
    """逐块产出 (delta_text, finish_reason, usage)；流结束时 finish_reason 非空，usage 可能为 None。"""
    stream = get_client().chat.completions.create(
        model=config.MODEL,
        max_tokens=config.MAX_TOKENS,
        messages=messages,
        stream=True,
        extra_headers={
            "HTTP-Referer": "http://localhost:5000",
            "X-Title": "CEO Interview Copilot",
        },
    )
    finish = None
    usage = None
    for chunk in stream:
        if getattr(chunk, "usage", None):
            usage = chunk.usage
        if not getattr(chunk, "choices", None):
            continue
        ch = chunk.choices[0]
        delta = getattr(ch.delta, "content", None) if getattr(ch, "delta", None) else None
        if delta:
            yield delta, None, None
        if getattr(ch, "finish_reason", None):
            finish = ch.finish_reason
    yield None, finish or "stop", usage


def sse_pack(obj):
    return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"


def stream_llm_sse(prompt, on_done=None):
    """以 SSE 流式返回 LLM 输出，持续吐数据避免线上代理超时；支持有限次自动续写。

    prompt 可以是字符串，也可以是一个无参函数（用于把耗时的预处理——如附件 OCR / Word 解析——
    延迟到流开始之后执行，期间持续发送心跳，避免连接在出首字节前被代理掐断）。
    注意：传入的函数不能依赖 Flask request 上下文（请在路由里先把所需数据读取好再闭包捕获）。
    """
    def gen():
        full = []
        # 立刻吐一个 SSE 注释，第一时间建立连接，预处理期间也保持连接活跃
        yield ": connected\n\n"
        try:
            if callable(prompt):
                box = {}
                done_evt = threading.Event()

                def _build():
                    try:
                        box["v"] = prompt()
                    except Exception as exc:
                        box["e"] = exc
                    finally:
                        done_evt.set()

                t = threading.Thread(target=_build, daemon=True)
                t.start()
                while not done_evt.wait(timeout=5):
                    yield ": working\n\n"
                if "e" in box:
                    raise box["e"]
                real_prompt = box["v"]
            else:
                real_prompt = prompt
            messages = [{"role": "user", "content": real_prompt}]
            max_cont = int(getattr(config, "MAX_CONTINUATIONS", 1))
            cont = 0
            print(f"[STREAM] start prompt_chars={len(real_prompt)} model={config.MODEL}")
            while True:
                finish = None
                last_usage = None
                for delta, fr, usage in _chat_stream(messages):
                    if delta:
                        full.append(delta)
                        yield sse_pack({"delta": delta})
                    if fr is not None:
                        finish = fr
                        last_usage = usage
                if last_usage:
                    log_usage(last_usage)
                if finish == "length" and cont < max_cont:
                    cont += 1
                    print(f"[STREAM] 截断，自动续写 第 {cont} 次")
                    content_so_far = "".join(full)
                    messages = [
                        {"role": "user", "content": prompt},
                        {"role": "assistant", "content": content_so_far},
                        {"role": "user", "content": (
                            "上一条回复因长度限制被截断了。请直接从被截断处继续输出剩余内容，"
                            "不要重复已经输出的部分，不要加任何前言或寒暄，不要重写开头。"
                        )},
                    ]
                    continue
                if finish == "length":
                    warn = "\n\n---\n> ⚠️ **内容较长，自动续写后仍未完成。可在 Render 调高 `MAX_TOKENS` 或 `MAX_CONTINUATIONS`。**"
                    full.append(warn)
                    yield sse_pack({"delta": warn})
                break
            content = "".join(full)
            meta = {"done": True}
            if on_done:
                try:
                    extra = on_done(content)
                    if extra:
                        meta.update(extra)
                except Exception as e:
                    print(f"[STREAM] on_done error: {e}")
            yield sse_pack(meta)
            print(f"[STREAM] done total_chars={len(content)}")
        except Exception as e:
            print(f"[STREAM ERROR] type={type(e).__name__} error={str(e)}")
            yield sse_pack({"error": friendly_llm_error(str(e))})
    return Response(
        stream_with_context(gen()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


def parse_pdf(file_stream):
    text_blocks = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_blocks.append(text)
    return "\n\n".join(text_blocks)


def parse_docx(file_stream):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 Word 解析依赖，请运行：pip3 install -r requirements.txt") from exc
    document = Document(file_stream)
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_image(file_stream):
    try:
        from PIL import Image
        import pytesseract
    except ImportError as exc:
        raise RuntimeError("缺少图片 OCR 依赖，请运行：pip3 install -r requirements.txt，并确保本机已安装 tesseract。") from exc
    image = Image.open(file_stream)
    return pytesseract.image_to_string(image, lang="chi_sim+eng")


def parse_interview_file_bytes(filename, data):
    """与 parse_interview_file 相同，但基于已读取的 bytes，可在无 request 上下文的线程中执行。"""
    filename = filename or "未命名文件"
    suffix = Path(filename).suffix.lower()
    stream = io.BytesIO(data)
    if suffix == ".pdf":
        text = parse_pdf(stream)
    elif suffix == ".docx":
        text = parse_docx(stream)
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        text = parse_image(stream)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="ignore")
    elif suffix == ".doc":
        raise RuntimeError(f"{filename} 是旧版 .doc 格式，暂不支持。请另存为 .docx 后上传。")
    else:
        raise RuntimeError(f"{filename} 的格式暂不支持，请上传 PDF、Word(.docx)、图片、TXT 或 Markdown。")
    if not text.strip():
        raise RuntimeError(f"{filename} 未解析出文字，可能是扫描件或图片 OCR 失败。")
    return f"## 文件：{filename}\n\n{text.strip()}"


def build_interview_notes(manual_notes, files_raw):
    """从已捕获的数据(无需 request)拼出面试纪要；files_raw 为 [(filename, bytes), ...]。"""
    parts = []
    if (manual_notes or "").strip():
        parts.append("## 手动输入\n\n" + manual_notes.strip())
    for filename, data in files_raw:
        parts.append(parse_interview_file_bytes(filename, data))
    return "\n\n---\n\n".join(parts)


def parse_interview_file(file_storage):
    filename = file_storage.filename or "未命名文件"
    suffix = Path(filename).suffix.lower()
    data = file_storage.read()
    stream = io.BytesIO(data)
    if suffix == ".pdf":
        text = parse_pdf(stream)
    elif suffix == ".docx":
        text = parse_docx(stream)
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        text = parse_image(stream)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="ignore")
    elif suffix == ".doc":
        raise RuntimeError(f"{filename} 是旧版 .doc 格式，暂不支持。请另存为 .docx 后上传。")
    else:
        raise RuntimeError(f"{filename} 的格式暂不支持，请上传 PDF、Word(.docx)、图片、TXT 或 Markdown。")
    if not text.strip():
        raise RuntimeError(f"{filename} 未解析出文字，可能是扫描件或图片 OCR 失败。")
    return f"## 文件：{filename}\n\n{text.strip()}"


def collect_interview_notes():
    manual_notes = ""
    file_notes = []
    if request.content_type and request.content_type.startswith("multipart/form-data"):
        manual_notes = request.form.get("interview_notes", "")
        for file_storage in request.files.getlist("interview_files"):
            if file_storage and file_storage.filename:
                file_notes.append(parse_interview_file(file_storage))
    else:
        data = request.get_json(silent=True) or {}
        manual_notes = data.get("interview_notes", "")
    parts = []
    if manual_notes.strip():
        parts.append("## 手动输入\n\n" + manual_notes.strip())
    parts.extend(file_notes)
    return "\n\n---\n\n".join(parts)


def render(template_name, variables):
    template = read_file(Path("prompts") / template_name)
    if not template:
        raise RuntimeError(f"找不到 Prompt 模板：prompts/{template_name}")
    for key, value in variables.items():
        template = template.replace("{" + key + "}", value or "（暂无）")
    return template


def save_output(candidate_name, filename, content):
    name = clean_person_name(candidate_name) or candidate_name or "候选人"
    safe_name = "".join(c for c in name if c not in r'\\/:*?"<>|').strip() or "候选人"
    date = current_date()
    folder = Path("outputs") / f"{date}_{safe_name}"
    folder.mkdir(parents=True, exist_ok=True)
    output_path = folder / filename
    output_path.write_text(content, encoding="utf-8")
    return str(output_path)


ARCHIVE_PATH = Path("outputs") / "archive.json"


def load_archive():
    if not ARCHIVE_PATH.exists():
        return []
    try:
        data = json.loads(ARCHIVE_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def save_archive(records):
    ARCHIVE_PATH.parent.mkdir(parents=True, exist_ok=True)
    tmp = ARCHIVE_PATH.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(ARCHIVE_PATH)


def _load_usage_records():
    if not USAGE_LOG_PATH.exists():
        return []
    try:
        data = json.loads(USAGE_LOG_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _summarize_by_date(records, days, source_filter=None):
    daily = defaultdict(lambda: {"calls": 0, "prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "cost": 0.0})
    cutoff = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
    for r in records:
        if r.get("date", "") < cutoff:
            continue
        if source_filter and r.get("source") != source_filter:
            continue
        d = r["date"]
        daily[d]["calls"] += 1
        daily[d]["prompt_tokens"] += r.get("prompt_tokens", 0)
        daily[d]["completion_tokens"] += r.get("completion_tokens", 0)
        daily[d]["total_tokens"] += r.get("total_tokens", 0)
        daily[d]["cost"] += r.get("cost", 0)
    summary = []
    for date in sorted(daily.keys(), reverse=True):
        summary.append({
            "date": date,
            "calls": daily[date]["calls"],
            "prompt_tokens": daily[date]["prompt_tokens"],
            "completion_tokens": daily[date]["completion_tokens"],
            "total_tokens": daily[date]["total_tokens"],
            "cost": round(daily[date]["cost"], 4),
        })
    total = {
        "calls": sum(d["calls"] for d in summary),
        "prompt_tokens": sum(d["prompt_tokens"] for d in summary),
        "completion_tokens": sum(d["completion_tokens"] for d in summary),
        "total_tokens": sum(d["total_tokens"] for d in summary),
        "cost": round(sum(d["cost"] for d in summary), 4),
    }
    return summary, total


def _merge_summaries(summaries_list):
    """把多组按天分组的 summary 合并成一组 combined。"""
    daily = defaultdict(lambda: {"calls": 0, "prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "cost": 0.0})
    for summary in summaries_list:
        for item in summary:
            d = daily[item["date"]]
            d["calls"] += item["calls"]
            d["prompt_tokens"] += item["prompt_tokens"]
            d["completion_tokens"] += item["completion_tokens"]
            d["total_tokens"] += item["total_tokens"]
            d["cost"] += item["cost"]
    merged = []
    for date in sorted(daily.keys(), reverse=True):
        merged.append({
            "date": date,
            "calls": daily[date]["calls"],
            "prompt_tokens": daily[date]["prompt_tokens"],
            "completion_tokens": daily[date]["completion_tokens"],
            "total_tokens": daily[date]["total_tokens"],
            "cost": round(daily[date]["cost"], 4),
        })
    total = {
        "calls": sum(d["calls"] for d in merged),
        "prompt_tokens": sum(d["prompt_tokens"] for d in merged),
        "completion_tokens": sum(d["completion_tokens"] for d in merged),
        "total_tokens": sum(d["total_tokens"] for d in merged),
        "cost": round(sum(d["cost"] for d in merged), 4),
    }
    return merged, total


@app.route("/api/usage", methods=["GET"])
@login_required
def api_usage():
    days = request.args.get("days", 7, type=int)
    days = max(1, min(days, 90))
    records = _load_usage_records()
    or_summary, or_total = _summarize_by_date(records, days, "openrouter")
    ide_summary, ide_total = _summarize_by_date(records, days, "ide")
    combined, combined_total = _merge_summaries([or_summary, ide_summary])
    return jsonify({
        "days": days,
        "openrouter": {"summary": or_summary, "total": or_total},
        "ide": {"summary": ide_summary, "total": ide_total},
        "combined": {"summary": combined, "total": combined_total},
    })


@app.route("/api/ide_usage", methods=["POST"])
@login_required
def api_ide_usage():
    """手动或通过脚本导入 IDE（Windsurf / Devin / Cursor 等）用量。
    Body JSON: {"date":"2026-06-24","prompt_tokens":1234,"completion_tokens":567,"cost":0.0123} 等
    """
    data = request.get_json(silent=True) or {}
    date_str = data.get("date") or datetime.now().strftime("%Y-%m-%d")
    prompt_tokens = int(data.get("prompt_tokens", 0))
    completion_tokens = int(data.get("completion_tokens", 0))
    total_tokens = int(data.get("total_tokens", prompt_tokens + completion_tokens))
    cost = float(data.get("cost", 0))
    try:
        records = _load_usage_records()
        records.append({
            "timestamp": datetime.now().isoformat(),
            "date": date_str,
            "source": "ide",
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
            "total_tokens": total_tokens,
            "cost": round(cost, 6),
            "model": data.get("model", "unknown"),
            "note": data.get("note", ""),
        })
        USAGE_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        tmp = USAGE_LOG_PATH.with_suffix(".json.tmp")
        tmp.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(USAGE_LOG_PATH)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/archive", methods=["GET"])
@login_required
def api_archive_get():
    return jsonify({"records": load_archive()})


@app.route("/api/archive", methods=["POST"])
@login_required
def api_archive_post():
    data = request.get_json(silent=True) or {}
    records = data.get("records")
    if not isinstance(records, list):
        return jsonify({"error": "records 必须是数组"}), 400
    try:
        save_archive(records)
        return jsonify({"ok": True, "count": len(records)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/")
@login_required
def index():
    return send_from_directory(".", "index.html")


@app.route("/assets/<path:filename>")
@login_required
def assets(filename):
    return send_from_directory("assets", filename)


@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "model": config.MODEL, "current_date": current_date()})


@app.route("/api/generate_jd", methods=["POST"])
@login_required
def api_generate_jd():
    data = request.get_json(silent=True) or {}
    raw_description = data.get("raw_description", "")
    if not raw_description.strip():
        return jsonify({"error": "岗位描述不能为空"}), 400
    return stream_llm_sse(jd_prompt(raw_description))


@app.route("/api/parse_resume", methods=["POST"])
@login_required
def api_parse_resume():
    if "file" not in request.files:
        return jsonify({"error": "未上传文件"}), 400
    file = request.files["file"]
    candidate_name = request.form.get("candidate_name", "候选人")
    try:
        raw_text = parse_pdf(file.stream)
        if not raw_text.strip():
            return jsonify({"error": "PDF 解析为空，可能是扫描件或图片型 PDF。"}), 400
        prompt = render("00_resume_parse.md", {
            "RAW_TEXT": raw_text,
            "候选人姓名": candidate_name,
            "CURRENT_DATE": current_date(),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    def on_done(content):
        extracted_name = extract_candidate_name(content)
        save_output(extracted_name, "01_简历.md", content)
        return {"candidate_name": extracted_name}
    return stream_llm_sse(prompt, on_done)


@app.route("/api/pre_interview", methods=["POST"])
@login_required
def api_pre_interview():
    data = request.get_json(silent=True) or {}
    candidate_name = data.get("candidate_name", "候选人")
    resume_md = data.get("resume_md", "")
    jd = data.get("jd", "")
    weights = data.get("weights", "")
    department_style = data.get("department_style", "")
    business_challenge = data.get("business_challenge", "")
    if not resume_md or not jd:
        return jsonify({"error": "简历或 JD 不能为空"}), 400
    try:
        prompt = render("01_pre_interview.md", {
            "JD": jd,
            "RESUME": resume_md,
            "AMIRO_STANDARD": read_file("standards/AMIRO_talent_standard.md"),
            "POSITION_WEIGHTS": weights,
            "DEPARTMENT_STYLE": department_style,
            "BUSINESS_CHALLENGE": business_challenge,
            "候选人姓名": candidate_name,
            "CURRENT_DATE": current_date(),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return stream_llm_sse(prompt, lambda content: save_output(candidate_name, "02_面试前作战图.md", content))


@app.route("/api/interview_plan", methods=["POST"])
@login_required
def api_interview_plan():
    data = request.get_json(silent=True) or {}
    candidate_name = data.get("candidate_name", "候选人")
    resume_md = data.get("resume_md", "")
    jd = data.get("jd", "")
    weights = data.get("weights", "")
    department_style = data.get("department_style", "")
    business_challenge = data.get("business_challenge", "")
    pre_report = data.get("pre_report", "")
    if not resume_md or not pre_report:
        return jsonify({"error": "简历或作战图不能为空"}), 400
    try:
        prompt = render("02_interview_plan.md", {
            "JD": jd,
            "RESUME": resume_md,
            "AMIRO_STANDARD": read_file("standards/AMIRO_talent_standard.md"),
            "POSITION_WEIGHTS": weights,
            "DEPARTMENT_STYLE": department_style,
            "BUSINESS_CHALLENGE": business_challenge,
            "PRE_INTERVIEW_REPORT": pre_report,
            "候选人姓名": candidate_name,
            "CURRENT_DATE": current_date(),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return stream_llm_sse(prompt, lambda content: save_output(candidate_name, "03_面试追问卡.md", content))


@app.route("/api/pre_screen_compare", methods=["POST"])
@login_required
def api_pre_screen_compare():
    data = request.get_json(silent=True) or {}
    position = (data.get("position") or "未分类岗位").strip()
    jd = data.get("jd", "")
    weights = data.get("weights", "")
    candidates = data.get("candidates", [])
    if not isinstance(candidates, list) or len(candidates) < 2:
        return jsonify({"error": "至少需要 2 位完成前三轮的候选人才能预筛对比"}), 400

    # 总量预算：模型上限 ~200K token，中文约 1 token/字，保留余量给模板与输出。
    # 候选人材料整体控制在约 12 万字以内，按人数均分每人预算，避免人多时超限。
    TOTAL_CHAR_BUDGET = 120000
    per_cand = max(2000, TOTAL_CHAR_BUDGET // max(1, len(candidates)))
    # 预筛排序最看重「简历 / 作战图」，追问卡价值较低且最长，预算分配靠后。
    resume_cap = int(per_cand * 0.45)
    pre_cap = int(per_cand * 0.40)
    plan_cap = per_cand - resume_cap - pre_cap

    def clip(text, cap):
        text = text or "（无）"
        if len(text) <= cap:
            return text
        return text[:cap].rstrip() + f"\n\n…（内容较长，为对比已截断，剩余约 {len(text) - cap} 字省略）"

    blocks = []
    for i, c in enumerate(candidates, 1):
        name = (c.get("name") or f"候选人{i}").strip()
        resume = clip(c.get("resume", ""), resume_cap)
        pre = clip(c.get("pre", ""), pre_cap)
        plan = clip(c.get("plan", ""), plan_cap)
        blocks.append(
            f"========== 候选人 {i}：{name} ==========\n\n"
            f"### 【{name}】简历解析\n{resume}\n\n"
            f"### 【{name}】面试前作战图\n{pre}\n\n"
            f"### 【{name}】面试追问卡（摘要）\n{plan}\n"
        )
    try:
        prompt = render("04_pre_screen_compare.md", {
            "JD": jd,
            "POSITION_WEIGHTS": weights,
            "AMIRO_STANDARD": read_file("standards/AMIRO_talent_standard.md"),
            "CANDIDATES": "\n\n".join(blocks),
            "岗位名": position,
            "CURRENT_DATE": current_date(),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return stream_llm_sse(prompt)


@app.route("/api/post_interview_eval", methods=["POST"])
@login_required
def api_post_interview_eval():
    data = request.form if request.content_type and request.content_type.startswith("multipart/form-data") else (request.get_json(silent=True) or {})
    candidate_name = data.get("candidate_name", "候选人")
    resume_md = data.get("resume_md", "")
    jd = data.get("jd", "")
    weights = data.get("weights", "")
    department_style = data.get("department_style", "")
    business_challenge = data.get("business_challenge", "")
    pre_report = data.get("pre_report", "")
    interview_plan = data.get("interview_plan", "")
    if not resume_md or not pre_report or not interview_plan:
        return jsonify({"error": "简历、作战图或追问卡不能为空"}), 400

    # 在请求上下文里快速读取上传附件的原始字节(I/O 很快)，把耗时的解析/OCR 延迟到流内进行
    manual_notes = ""
    files_raw = []
    if request.content_type and request.content_type.startswith("multipart/form-data"):
        manual_notes = request.form.get("interview_notes", "")
        for fs in request.files.getlist("interview_files"):
            if fs and fs.filename:
                files_raw.append((fs.filename, fs.read()))
    else:
        manual_notes = (request.get_json(silent=True) or {}).get("interview_notes", "")
    if not manual_notes.strip() and not files_raw:
        return jsonify({"error": "缺少面试记录 / 附件，无法生成评判"}), 400

    cur_date = current_date()
    amiro_standard = read_file("standards/AMIRO_talent_standard.md")

    def build_prompt():
        interview_notes = build_interview_notes(manual_notes, files_raw)
        if not interview_notes.strip():
            raise RuntimeError("面试记录解析为空，请检查附件内容或改用文字粘贴。")
        return render("03_post_interview_eval.md", {
            "JD": jd,
            "RESUME": resume_md,
            "AMIRO_STANDARD": amiro_standard,
            "POSITION_WEIGHTS": weights,
            "DEPARTMENT_STYLE": department_style,
            "BUSINESS_CHALLENGE": business_challenge,
            "PRE_INTERVIEW_REPORT": pre_report,
            "INTERVIEW_PLAN": interview_plan,
            "INTERVIEW_NOTES": interview_notes,
            "候选人姓名": candidate_name,
            "CURRENT_DATE": cur_date,
        })
    return stream_llm_sse(build_prompt, lambda content: save_output(candidate_name, "04_面试后评判.md", content))


if __name__ == "__main__":
    import threading
    import time
    import webbrowser

    port = getattr(config, "PORT", 5000)
    if not getattr(config, "API_KEY", ""):
        print_error("当前没有设置 OPENROUTER_API_KEY，页面可以打开，但 AI 生成功能会提示缺少 Key。")

    def open_browser():
        time.sleep(1.5)
        webbrowser.open(f"http://localhost:{port}")

    threading.Thread(target=open_browser, daemon=True).start()
    print("\n" + "=" * 60)
    print("  🎯 CEO 面试副驾驶 - 本地版已启动")
    print(f"  访问地址: http://localhost:{port}")
    print(f"  使用模型: {config.MODEL}")
    print("=" * 60 + "\n")
    app.run(host="127.0.0.1", port=port, debug=False)
